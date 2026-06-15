import os
import re


def extract_text_from_pdf(filepath):
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(filepath):
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text_from_txt(filepath):
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Failed to read TXT file: {str(e)}")


def parse_resume(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        raw_text = extract_text_from_pdf(filepath)
    elif ext == ".docx":
        raw_text = extract_text_from_docx(filepath)
    elif ext == ".txt":
        raw_text = extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")
    return normalize_text(raw_text)


def normalize_text(text):
    # Collapse whitespace
    text = re.sub(r'\r\n|\r', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
